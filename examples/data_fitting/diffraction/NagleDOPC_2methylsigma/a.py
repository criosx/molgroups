import docx
import os
import sys
from docx.shared import Inches
import pandas as pd
from IPython.display import display
import math

sys.path.append(os.getcwd())


document = docx.Document()


def addfit(name):
    path = os.getcwd() + "/" + name + "/T3/"
    
    if "new" in name:
        document.add_heading("DOPC (2005)", 0)
    else:
        document.add_heading("DOPC (stitched)", 0)

    document.add_heading("Fit", level=1)
    document.add_picture(path + "run-model" + ".png")

    document.add_heading("Parameters", level=1)

    def rounds(n, s=4):
        return round(n, s - int(math.floor(math.log10(abs(n)))) - 1)

    df = pd.read_json(path + "run.json").T
    cols = ["label", "best", "mean", "median", "p68"]
    for elem in cols:
        if elem == "label":
            continue
        for i, item in enumerate(df[elem]):
            if isinstance(item, list):
                df[elem][i] = [rounds(i, 4) for i in item]
            else:
                df[elem][i] = rounds(item, 4)

    # display(test)
    test = df[cols].values.tolist()

    table = document.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells

    for i, elem in enumerate(cols):
        hdr_cells[i].text = elem
    
    hdr_cells[0].text = ""
    hdr_cells[1].text = "fit range"
    hdr_cells[-1].text = "68% confidence interval"
    for label, best, mean, median, p68 in test:
        row_cells = table.add_row().cells
        row_cells[0].text = str(label)
        row_cells[2].text = str(best)
        row_cells[3].text = str(mean)
        row_cells[4].text = str(median)
        row_cells[5].text = str(p68)
    
    # document.add_heading("variable histogram", level=1)
    # document.add_picture(path + "var" + ".png", width=Inches(6))

    # document.add_heading("correlation matrix", level=1)
    # document.add_picture(path + "corr" + ".png")

    document.add_heading("Statistics (median, sigma confidence interval)", level=1)

    text = ""
    with open(path+"CalculationResults.dat") as fp:
        for i, line in enumerate(fp):
            if i == 12:
                text = 'Area per lipid (Å^2) [' + line.split('[')[1] + text
                break
    document.add_paragraph(text)

    document.add_heading("Electron Density Profile", level=1)
    document.add_picture(path + "electron_density" + ".png")

    document.add_heading("Component Volume Occupancy", level=1)
    document.add_picture(path + "cvo" + ".png")

names = ["new", "dopc"]
for name in names:
    addfit(name)
    document.add_page_break()

document.save("DOPC.docx")
